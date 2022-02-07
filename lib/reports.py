from lib.models import session, Disk, Partition, Fstab, Directory, Pool
from lib.parsers import hosts
import os
from datetime import datetime
import openpyxl
from openpyxl.styles import PatternFill
from openpyxl.utils.cell import get_column_letter
import docx
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENTATION

class Report():
    def __init__(self):
        self.hosts = hosts()
        if not os.path.isdir('reports'):
            os.mkdir('reports')

    def inventory(self):
        book = openpyxl.Workbook()
        sheet = book.active
        sheet.title = "Инвентаризация жестких дисков"
        head = ('Сервер', 'Имя', 'Объем', 'Описание', 'Производитель', 'Продукт', 'Серийный №')
        sheet.append(head)
        for col in range(len(head)):
            sheet.cell(1, col+1).fill = PatternFill(fgColor="00FF00", fill_type="solid")
            sheet.column_dimensions[get_column_letter(col+1)].width = 20
        row = 2
        equipment = session.query(Disk).all()
        for disk in equipment:
            sheet.cell(row=row, column=1).value = self.hosts[disk.ip]
            col = 2
            for param in ('name', 'size', 'description', 'vendor', 'product', 'serial'):
                sheet.cell(row=row, column=col).value = getattr(disk, param)
                col += 1
            row += 1
        sheet.auto_filter.add_sort_condition('A:A')
        book.save("reports/Инвентаризация дисков.xlsx")

    def bugs(self):
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)
        # Landscape orientation, A4 paper
        section = doc.sections[-1]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        # Margin
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        # Header
        d = datetime.now()
        section.header.add_paragraph(f'Отчет сформирован {d.year}.{d.month}.{d.day} в {d.hour}:{d.minute}')
        for ip in self.hosts:
            doc.add_heading(f"{self.hosts[ip].upper()} ({ip})", 0)
            disks = session.query(Disk).filter(Disk.ip==ip).count()
            if not disks:
                doc.add_heading(f"Не найдены результаты сканирования для {ip}", 5)
                continue
            partitions = session.query(Partition).outerjoin(Disk).filter(Disk.ip==ip).all()
            for part in partitions:
                # Partitions with plots has size few TB and must be mounted
                if part.size[-1]=='T' and part.fstab==0:
                    doc.add_paragraph(f"{part.name} не имеет записи в fstab")
                if part.fstab > 1:
                    doc.add_paragraph(f"{part.name} имеет {part.fstab} записей в fstab")
            fstab_errors = session.query(Fstab).filter(Fstab.ip==ip, Fstab.error.isnot(None)).count()
            if not fstab_errors:
                doc.add_heading("Ошибок fstab не обнаружено", 6)
            else:
                doc.add_heading("Ошибки fstab:", 6)
                fstab = session.query(Fstab).filter(Fstab.ip==ip).all()
                table = doc.add_table(rows=len(fstab), cols=3)
                for s in fstab:
                    i = s.sn - 1
                    cell = table.cell(i, 0)
                    cell.text = f"{s.sn}."
                    cell = table.cell(i, 1)
                    cell.text = s.ss
                    cell = table.cell(i, 2)
                    cell.text = s.error or ''
                table.columns[0].width = Cm(1.0)
                table.columns[1].width = Cm(17.0)
                table.columns[2].width = Cm(8.0)
            plot_dirs = session.query(Directory).filter(Directory.ip==ip).count()
            if not plot_dirs:
                doc.add_heading("Плоты chia не найдены", 6)
            else:
                unused = session.query(Directory).filter(Directory.ip==ip, Directory.used==0).all()
                double = session.query(Directory).filter(Directory.ip==ip, Directory.used>1).all()
                if unused:
                    doc.add_heading("Неиспользуемые плоты chia:", 6)
                    for dir in unused:
                        doc.add_paragraph(dir.path)
                if double:
                    doc.add_heading("Плоты chia, используемые двумя фармерами", 6)
                    for dir in unused:
                        doc.add_paragraph(dir.path)
                if not unused and not double:
                    doc.add_heading("Ошибок в конфигурации фармера не обнаружено", 6)
            pool_errors = session.query(Pool).filter(Pool.ip==ip, Pool.errors.isnot(None)).all()
            if not pool_errors:
                doc.add_heading("Неиспользуемых или используемых дважды плотов не обнаружено", 6)
            else:
                hpool = session.query(Pool).filter(Pool.ip==ip, Pool.pool=='hpool').one_or_none()
                if hpool:
                    if hpool.errors:
                        doc.add_heading("Найдены плоты, но не используются фармером hpool:", 6)
                        doc.add_paragraph(hpool.errors)
                    else:
                        doc.add_heading("Фермер hpool сконфигурирован правильно")
                flex = session.query(Pool).filter(Pool.ip==ip, Pool.pool=='flexpool').one_or_none()
                if flex:
                    if flex.errors:
                        doc.add_heading("Найдены плоты, но не используются фармером flexpool:")
                        doc.add_paragraph(flex.errors)
                    else:
                        doc.add_heading("Фермер flexpool сконфигурирован правильно")
        doc.save('reports/Ошибки конфигурации.docx')

    def usage(self):
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)
        section = doc.sections[-1]
        # Header
        d = datetime.now()
        section.header.add_paragraph(f'Отчет сформирован {d.year}.{d.month}.{d.day} в {d.hour}:{d.minute}')
        for ip in self.hosts:
            doc.add_heading(f"{self.hosts[ip].upper()} ({ip})", 0)
            partitions = session.query(Partition).outerjoin(Disk).filter(Disk.ip==ip).all()
            hdr = ('Имя', 'Точка монтирования', 'Объем', 'Использовано', 'Доступно', '%')
            table = doc.add_table(rows=len(partitions)+1, cols=len(hdr))
            for j, h in enumerate(hdr):
                cell = table.cell(0, j)
                cell.text = h
                cell.paragraphs[0].runs[0].font.bold = True
            for i, part in enumerate(partitions):
                for j, attr in enumerate(('name', 'mountpoint', 'size', 'used', 'available')):
                    cell = table.cell(i+1, j)
                    cell.text = getattr(part, attr) or ''
                cell = table.cell(i+1, 5)
                cell.text = f"{part.usage}%" if part.usage else ''
        doc.save('reports/Использование дискового пространства.docx')

    @property
    def cache_exists(self):
        data = session.query(Disk).count()
        if data:
            return True
        return False